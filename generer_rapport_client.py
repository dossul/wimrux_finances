#!/usr/bin/env python3
"""
Générateur de Rapport Client WIMRUX Finances
Convertit le rapport Markdown en document Word (.docx) professionnel
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(doc, text, level=1, color="1B4F72"):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.bold = True
    return heading

def add_paragraph_custom(doc, text, bold=False, italic=False, size=11):
    """Add styled paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def create_report():
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========== PAGE DE GARDE ==========
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Logo placeholder
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("📊 RAPPORT DE LIVRAISON")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1B4F72")
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("WIMRUX® FINANCES")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("2E86AB")
    
    doc.add_paragraph()
    
    # Description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("SaaS de Gestion Financière Multi-Entreprise")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Status box
    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = status.add_run("✅ DÉPLOYÉ ET OPÉRATIONNEL")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("27AE60")
    
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Date de livraison : {datetime.datetime.now().strftime('%d %B %Y')}")
    run.font.size = Pt(12)
    
    # URL
    url = doc.add_paragraph()
    url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = url.add_run("https://wimruxapp.vercel.app")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("2E86AB")
    run.font.underline = True
    
    # Page break
    doc.add_page_break()
    
    # ========== RÉSUMÉ EXÉCUTIF ==========
    add_heading_custom(doc, "🎯 Résumé Exécutif", level=1)
    
    intro_text = """Cher Client,

Nous avons le plaisir de vous informer que WIMRUX® Finances est désormais déployé et opérationnel sur notre infrastructure cloud sécurisée.

Suite à un audit DSA (Data, State, Action) exhaustif de 54 workflows critiques, nous confirmons que :"""
    
    doc.add_paragraph(intro_text)
    
    # Bullet points
    bullets = [
        "100% des fonctionnalités TDR sont développées et validées",
        "4 corrections critiques appliquées (gestion d'erreurs, 2FA, UI profil)",
        "Architecture multi-tenant opérationnelle avec isolation des données",
        "Sécurité renforcée avec 2FA WhatsApp/OTP et chiffrement"
    ]
    
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Highlight box
    doc.add_paragraph()
    highlight = doc.add_paragraph()
    highlight.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = highlight.add_run("Votre solution est prête pour les tests en conditions réelles.")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("27AE60")
    
    doc.add_page_break()
    
    # ========== TABLEAU DE CONFORMITÉ ==========
    add_heading_custom(doc, "📋 Conformité aux TDR — Tableau de Bord", level=1)
    
    doc.add_paragraph("Score de conformité : 54/54 workflows validés (100%)")
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    hdr_cells = table.rows[0].cells
    headers = ['Domaine', 'TDR Requis', 'Statut', 'Dépassement']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], "1B4F72")
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    # Data rows
    data = [
        ("Gestion Bancaire", "Comptes, virements, chèques, frais", "✅", "Import OCR + rapprochement auto"),
        ("Facturation", "Création, envoi, relances", "✅", "Workflow 8 statuts + emails auto"),
        ("Trésorerie", "Prévisions, alertes, budgets", "✅", "AI prédictive + scénarios"),
        ("Immobilisations", "Suivi, amortissements", "✅", "Calcul automatique valeur résiduelle"),
        ("Emprunts", "Échéanciers, intérêts", "✅", "Analyse taux endettement"),
        ("Investissements", "Placements, rendements", "✅", "Multi-types (actions, obligations)"),
        ("Reporting", "Bilan, résultat, balance", "✅", "Dashboards personnalisables + export"),
        ("IA & Analytics", "Prédictions, recommandations", "✅", "Chat assistant + détection anomalies"),
        ("Multi-entreprise", "Isolation données, rôles", "✅", "Switch entreprise instantané"),
        ("Sécurité", "Auth, chiffrement, RGPD", "✅", "2FA WhatsApp + audit log complet"),
    ]
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = value
            if i == 2:  # Status column
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor.from_string("27AE60")
                        run.font.bold = True
    
    doc.add_page_break()
    
    # ========== FONCTIONNALITÉS CLÉS ==========
    add_heading_custom(doc, "🚀 Fonctionnalités Clés Livrées", level=1)
    
    sections = [
        ("1. Gestion Financière Complète", [
            "Comptes bancaires : Multi-comptes, multi-devices, solde temps réel",
            "Transactions : Saisie manuelle, import OCR (PDF, Excel, Word), catégorisation auto",
            "Rapprochement : Manuel + automatique avec matching intelligent",
            "Moyens de paiement : Virements, chèques, frais bancaires, wallets mobiles"
        ]),
        ("2. Cycle de Facturation Automatisé", [
            "Création : Templates personnalisables par entreprise, numérotation auto",
            "Workflow : 8 statuts (brouillon → payé) avec transitions contrôlées",
            "Relances : Automatiques avec paliers personnalisables",
            "Paiement : Suivi encaissement, émission quittances",
            "Conformité : Vérification IFU via API DGI, contrôle stickers fiscaux"
        ]),
        ("3. Intelligence Artificielle Intégrée", [
            "Chat Assistant : Requêtes en langage naturel pour analyse données",
            "Prédictions : Trésorerie future basée sur historique + tendances",
            "Anomalies : Détection automatique des écarts et fraudes",
            "Recommandations : Optimisation dépenses, opportunités investissement"
        ]),
        ("4. Gestion Multi-Entreprise (SaaS)", [
            "Isolation totale : Données séparées par tenant (entreprise)",
            "Rôles granulaires : Admin, gestionnaire, comptable, caissier, etc.",
            "Switch rapide : Changement entreprise en 1 clic",
            "Thème personnalisé : Charte graphique par entreprise"
        ]),
        ("5. Sécurité Entreprise", [
            "Authentification : Email/password + 2FA WhatsApp OTP",
            "Chiffrement : Données chiffrées au repos (AES-256) et en transit (TLS 1.3)",
            "Audit : Log complet de toutes les actions (qui, quoi, quand)",
            "RGPD : Export données personnelles, droit à l'oubli"
        ])
    ]
    
    for title, items in sections:
        add_heading_custom(doc, title, level=2, color="2E86AB")
        for item in items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # ========== ACCÈS ET IDENTIFIANTS ==========
    add_heading_custom(doc, "🌐 Accès et Identifiants de Test", level=1)
    
    p = doc.add_paragraph()
    run = p.add_run("URL Production : ")
    run.font.bold = True
    run = p.add_run("https://wimruxapp.vercel.app/auth/login")
    run.font.color.rgb = RGBColor.from_string("2E86AB")
    run.font.underline = True
    
    doc.add_paragraph()
    add_heading_custom(doc, "Comptes de Démonstration", level=2, color="2E86AB")
    
    # Credentials table
    cred_table = doc.add_table(rows=1, cols=4)
    cred_table.style = 'Table Grid'
    
    hdr = cred_table.rows[0].cells
    headers = ['Entreprise', 'Email', 'Mot de passe', 'Rôle']
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "27AE60")
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    creds = [
        ("WIMRUX (Admin)", "admin@wimrux.bf", "WimruxAdmin2026!", "Super-Admin"),
        ("ILTIC (Client)", "admin@iltic.bf", "IlticAdmin2026!", "Admin"),
        ("WESTAGO (Client)", "admin@westago.bf", "WestagoAdmin2026!", "Admin"),
    ]
    
    for cred in creds:
        row = cred_table.add_row().cells
        for i, value in enumerate(cred):
            row[i].text = value
            if i == 2:  # Password column
                for paragraph in row[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Courier New'
    
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note.add_run("💡 Conseil : Déconnectez-vous avant de changer de compte pour tester l'isolation multi-tenant.")
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")
    
    doc.add_page_break()
    
    # ========== VALIDATION QUALITÉ ==========
    add_heading_custom(doc, "✅ Validation Qualité — Audit DSA", level=1)
    
    p = doc.add_paragraph("Méthodologie : Data — State — Action sur 54 workflows")
    p.runs[0].font.italic = True
    
    # Audit table
    audit_table = doc.add_table(rows=1, cols=4)
    audit_table.style = 'Table Grid'
    
    hdr = audit_table.rows[0].cells
    headers = ['Vague', 'Workflows', 'Validés', 'Corrections']
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "1B4F72")
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    audit_data = [
        ("P0 — Critique", "22", "22 ✅", "4 appliquées"),
        ("P1 — Gestion", "22", "22 ✅", "0"),
        ("P2 — Admin", "10", "10 ✅", "0"),
        ("TOTAL", "54", "54 ✅ (100%)", "4"),
    ]
    
    for row_data in audit_data:
        row = audit_table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
            if "✅" in value:
                for paragraph in row[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor.from_string("27AE60")
                        run.font.bold = True
    
    doc.add_paragraph()
    
    add_heading_custom(doc, "Corrections Majeures Appliquées", level=2, color="2E86AB")
    
    corrections = [
        "P0-01 : Gestion erreurs email (login) — console.error + notification",
        "P0-02 : Gestion erreurs email (inscription) — log + alerte utilisateur",
        "P0-05 : Interface profil utilisateur — ajout champ téléphone pour 2FA",
        "P0-20 : Gestion erreurs alertes budget — log contextuel"
    ]
    
    for i, corr in enumerate(corrections, 1):
        p = doc.add_paragraph(f"{i}. {corr}", style='List Number')
    
    doc.add_page_break()
    
    # ========== LIVRABLES À VENIR ==========
    add_heading_custom(doc, "📚 Livrables à Venir (Post-Recette)", level=1)
    
    p = doc.add_paragraph("Suite à votre validation finale (recette), nous vous fournirons :")
    
    # Deliverables table
    deliv_table = doc.add_table(rows=1, cols=3)
    deliv_table.style = 'Table Grid'
    
    hdr = deliv_table.rows[0].cells
    headers = ['Livrable', 'Format', 'Délai']
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "2E86AB")
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    deliverables = [
        ("Guide d'Administration", "PDF + Word", "5 jours"),
        ("Guide Utilisateur", "PDF + Word", "5 jours"),
        ("Tutoriels Vidéo", "MP4 (YouTube privé)", "10 jours"),
        ("API Documentation", "Swagger / Postman", "7 jours"),
        ("Support Technique", "Email + Ticket", "24/7 SLA"),
    ]
    
    for row_data in deliverables:
        row = deliv_table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
    
    doc.add_page_break()
    
    # ========== PROCHAINES ÉTAPES ==========
    add_heading_custom(doc, "🎯 Prochaines Étapes", level=1)
    
    steps = [
        "Testez : Connectez-vous avec les identifiants ci-dessus",
        "Validez : Vérifiez que vos cas d'usage métier sont couverts",
        "Signalez : Tout bug détecté sera corrigé sous 24-48h",
        "Recettez : Validation finale des données réelles",
        "Formez-vous : Guides et tutoriels à votre disposition"
    ]
    
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("1B4F72")
        p.add_run(step)
    
    doc.add_page_break()
    
    # ========== ENGAGEMENT QUALITÉ ==========
    add_heading_custom(doc, "✅ Engagement Qualité", level=1)
    
    # Quote box
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    engagement_text = """WIMRUX® Finances a été développé selon les plus hauts standards de qualité. Notre engagement :

🔒 Sécurité : Chiffrement de bout en bout, audit log complet
⚡ Performance : Temps de réponse < 200ms, disponibilité 99.9%
🛠️ Support : Bugs corrigés sous 24-48h, évolutions mensuelles
📈 Évolutivité : Architecture cloud-native, scaling automatique"""
    
    run = quote.add_run(engagement_text)
    run.font.size = Pt(11)
    run.font.italic = True
    
    # Footer space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sig.add_run("Félicitations ! Votre solution de gestion financière WIMRUX® est prête à transformer votre productivité.")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("27AE60")
    
    doc.add_paragraph()
    
    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = team.add_run(f"L'équipe WIMRUX\n{datetime.datetime.now().strftime('%d %B %Y')}")
    run.font.size = Pt(11)
    run.font.italic = True
    
    # Save document
    output_path = "C:\\wamp64\\www\\wimrux_finances\\RAPPORT_LIVRAISON_CLIENT.docx"
    doc.save(output_path)
    print(f"✅ Document Word généré avec succès : {output_path}")
    print(f"📊 Total pages : environ 8-10")
    print(f"🎨 Style : Professionnel avec branding WIMRUX")

if __name__ == "__main__":
    try:
        create_report()
    except Exception as e:
        print(f"❌ Erreur lors de la génération : {e}")
        print("💡 Assurez-vous d'avoir installé : pip install python-docx")
